"""Main window — Balabolka-style testing and settings UI.

Mirrors the Rust/Tauri frontend design:
  Toolbar → Engine tabs → Voice controls bar → Audio sliders → Text editor → Status bar

Opened from the tray "Settings…" item (or double-click in the future).

This is a reactive View in the MVC split (see app_state.py/app_controller.py):
it holds no state of its own beyond widget contents, reads initial values
from AppState, connects to AppState's signals to stay in sync (including
changes that originate elsewhere — the tray menu, Load Settings, ...), and
calls AppController methods in response to user input. It never mutates
AppState directly.
"""
from __future__ import annotations

from pathlib import Path

from PySide6.QtCore import QPoint, QSize, Qt, QTimer, Signal
from PySide6.QtGui import (
    QAction,
    QColor,
    QFont,
    QIcon,
    QPainter,
    QPixmap,
    QPolygon,
)
from PySide6.QtWidgets import (
    QComboBox,
    QFileDialog,
    QFrame,
    QHBoxLayout,
    QLabel,
    QMainWindow,
    QMenuBar,
    QMessageBox,
    QPushButton,
    QSizePolicy,
    QSlider,
    QStatusBar,
    QTabWidget,
    QTextEdit,
    QToolBar,
    QToolButton,
    QVBoxLayout,
    QWidget,
)

from ..control.app_controller import AppController
from ..model.app_state import AppState
from ..config import get_controls, get_voices, load_effective_config
from ..device import cuda_available
from ..engines.registry import StackInfo

_ACCENT = "#0078d4"
_TOOLBAR_BG = "#f5f5f5"
_TABS_BG = "#e8e8e8"
_VOICEBAR_BG = "#f5f5f5"
_SLIDER_BG = "#fafafa"


class _SliderRow(QWidget):
    """Label + horizontal slider + live value — horizontal layout for the slider strip."""

    valueChanged = Signal(float)

    def __init__(
        self,
        label: str,
        min_val: float,
        max_val: float,
        default: float,
        step: float = 1.0,
        fmt: str = "{:.0f}",
        enabled: bool = True,
        parent: QWidget | None = None,
    ) -> None:
        super().__init__(parent)
        self._fmt = fmt
        self._step = step
        self._min = min_val
        self._max = max_val

        vbox = QVBoxLayout(self)
        vbox.setContentsMargins(0, 0, 0, 0)
        vbox.setSpacing(4)

        top = QHBoxLayout()
        top.setContentsMargins(0, 0, 0, 0)
        self._name_lbl = QLabel(label)
        self._name_lbl.setStyleSheet("font-size:11px; color:#666;")
        self._val_lbl = QLabel(fmt.format(default))
        self._val_lbl.setStyleSheet("font-size:11px; font-weight:600; color:#333;")
        top.addWidget(self._name_lbl)
        top.addStretch()
        top.addWidget(self._val_lbl)
        vbox.addLayout(top)

        self._slider = QSlider(Qt.Orientation.Horizontal)
        # Qt sliders are integer-only; scale by 1000 for float precision
        scale = 1000
        self._scale = scale
        self._slider.setRange(int(min_val * scale), int(max_val * scale))
        self._slider.setValue(int(default * scale))
        self._slider.setSingleStep(max(1, int(step * scale)))
        self._slider.setEnabled(enabled)
        self._slider.setStyleSheet(f"""
            QSlider::groove:horizontal {{
                height: 6px;
                background: #d0d0d0;
                border-radius: 3px;
            }}
            QSlider::sub-page:horizontal {{
                background: {_ACCENT};
                border-radius: 3px;
            }}
            QSlider::handle:horizontal {{
                width: 14px;
                height: 14px;
                margin: -4px 0;
                background: #ffffff;
                border: 2px solid {_ACCENT};
                border-radius: 7px;
            }}
            QSlider::handle:horizontal:disabled {{
                border-color: #aaa;
            }}
        """)
        self._slider.valueChanged.connect(self._on_change)
        vbox.addWidget(self._slider)

    def _on_change(self, raw: int) -> None:
        v = raw / self._scale
        self._val_lbl.setText(self._fmt.format(v))
        self.valueChanged.emit(v)

    def value(self) -> float:
        return self._slider.value() / self._scale

    def setValue(self, v: float) -> None:
        self._slider.setValue(int(v * self._scale))

    def setEnabled(self, enabled: bool) -> None:
        self._slider.setEnabled(enabled)
        self._name_lbl.setStyleSheet(
            f"font-size:11px; color:{'#666' if enabled else '#aaa'};"
        )


class _MultiFormatEditor(QWidget):
    """Text editor that can open and save TXT, MD, HTML, DOCX, and PDF files.

    The editor always stays visible regardless of which engine tab is active
    so content is never lost when switching stacks.
    """

    textChanged = Signal()

    _OPEN_FILTER = (
        "All Supported (*.txt *.md *.html *.htm *.docx *.pdf);;"
        "Plain Text (*.txt);;"
        "Markdown (*.md);;"
        "HTML (*.html *.htm);;"
        "Word Document (*.docx);;"
        "PDF (*.pdf)"
    )
    _SAVE_FILTER = (
        "Plain Text (*.txt);;"
        "Markdown (*.md);;"
        "HTML (*.html *.htm);;"
        "Word Document (*.docx)"
    )

    def __init__(self, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        self._current_path: Path | None = None
        self._current_fmt: str = "txt"

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        self._edit = QTextEdit()
        self._edit.setPlaceholderText("Enter or paste text here…")
        self._edit.setFont(QFont("Consolas", 11))
        self._edit.setAcceptRichText(True)
        self._edit.setStyleSheet(
            "QTextEdit { border: none; padding: 8px; background: #ffffff; }"
        )
        self._edit.textChanged.connect(self.textChanged)
        layout.addWidget(self._edit)

    # ── Public API used by MainWindow ─────────────────────────────────────────

    def to_plain_text(self) -> str:
        return self._edit.toPlainText()

    def clear(self) -> None:
        self._edit.clear()
        self._current_path = None
        self._current_fmt = "txt"

    def current_path(self) -> Path | None:
        return self._current_path

    def current_fmt(self) -> str:
        return self._current_fmt

    def open_file(self, start_dir: str = "") -> Path | None:
        """Show open dialog and load the chosen file. Returns path or None."""
        path_str, _ = QFileDialog.getOpenFileName(
            self, "Open", start_dir, self._OPEN_FILTER
        )
        if not path_str:
            return None
        path = Path(path_str)
        self._load(path)
        return path

    def save_file(self, force_dialog: bool = False) -> Path | None:
        """Save to current path (or show dialog if no path or force_dialog)."""
        if not force_dialog and self._current_path and self._current_fmt != "pdf":
            self._save(self._current_path)
            return self._current_path
        path_str, _ = QFileDialog.getSaveFileName(
            self, "Save As",
            str(self._current_path) if self._current_path else "",
            self._SAVE_FILTER,
        )
        if not path_str:
            return None
        path = Path(path_str)
        self._save(path)
        return path

    # ── Load ──────────────────────────────────────────────────────────────────

    def _load(self, path: Path) -> None:
        ext = path.suffix.lower()
        try:
            if ext in (".txt", ".md"):
                self._edit.setPlainText(path.read_text(encoding="utf-8", errors="replace"))
                self._current_fmt = ext.lstrip(".")
            elif ext in (".html", ".htm"):
                self._edit.setHtml(path.read_text(encoding="utf-8", errors="replace"))
                self._current_fmt = "html"
            elif ext == ".docx":
                self._edit.setPlainText(self._read_docx(path))
                self._current_fmt = "docx"
            elif ext == ".pdf":
                self._edit.setPlainText(self._read_pdf(path))
                self._current_fmt = "pdf"  # read-only; save → dialog forces TXT/MD/HTML
            else:
                self._edit.setPlainText(path.read_text(encoding="utf-8", errors="replace"))
                self._current_fmt = "txt"
            self._current_path = path
        except Exception as exc:
            QMessageBox.warning(self, "Open Failed", str(exc))

    def _read_docx(self, path: Path) -> str:
        try:
            from docx import Document  # type: ignore[import-untyped]
            doc = Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            raise RuntimeError(
                "python-docx is not installed.\nRun: pip install python-docx"
            ) from None

    def _read_pdf(self, path: Path) -> str:
        try:
            import fitz  # PyMuPDF  # type: ignore[import-untyped]
            doc = fitz.open(str(path))
            return "\n\n".join(page.get_text() for page in doc)
        except ImportError:
            pass
        try:
            import pdfplumber  # type: ignore[import-untyped]
            with pdfplumber.open(str(path)) as pdf:
                return "\n\n".join(page.extract_text() or "" for page in pdf.pages)
        except ImportError:
            raise RuntimeError(
                "PDF reading requires PyMuPDF or pdfplumber.\n"
                "Run: pip install pymupdf   or   pip install pdfplumber"
            ) from None

    # ── Save ──────────────────────────────────────────────────────────────────

    def _save(self, path: Path) -> None:
        ext = path.suffix.lower()
        try:
            if ext in (".html", ".htm"):
                path.write_text(self._edit.toHtml(), encoding="utf-8")
            elif ext == ".docx":
                self._write_docx(path)
            else:
                path.write_text(self._edit.toPlainText(), encoding="utf-8")
            self._current_path = path
            self._current_fmt = ext.lstrip(".") if ext != ".htm" else "html"
        except Exception as exc:
            QMessageBox.warning(self, "Save Failed", str(exc))

    def _write_docx(self, path: Path) -> None:
        try:
            from docx import Document  # type: ignore[import-untyped]
            doc = Document()
            for line in self._edit.toPlainText().split("\n"):
                doc.add_paragraph(line)
            doc.save(str(path))
        except ImportError:
            raise RuntimeError(
                "python-docx is not installed.\nRun: pip install python-docx"
            ) from None


class MainWindow(QMainWindow):
    """Balabolka-style main window — testing harness and settings.

    Reactive View: reads AppState, connects to its signals, calls
    AppController methods. Holds no state of its own beyond widget
    contents (and those are always kept in sync with AppState, never a
    separate source of truth).
    """

    def __init__(
        self,
        app_state: AppState,
        controller: AppController,
        models_root: Path | None = None,
        parent: QWidget | None = None,
    ) -> None:
        super().__init__(parent)
        self.setWindowTitle("AlienVox")
        self.resize(720, 540)
        self.setMinimumSize(560, 400)
        # Qt.Window ensures it gets its own taskbar button even when parented
        self.setWindowFlags(Qt.WindowType.Window)
        _icon_path = Path(__file__).parent / "resources" / "icons" / "icon_256x256.png"
        if _icon_path.exists():
            self.setWindowIcon(QIcon(str(_icon_path)))

        self._state = app_state
        self._controller = controller
        self._models_root = models_root

        # Collectors for post-build wiring
        self._voice_combos: list[tuple[QComboBox, str]] = []
        self._model_combos: list[tuple[QComboBox, str]] = []
        self._sliders: list[tuple[str, _SliderRow]] = []

        # Slider debounce timer
        self._slider_save_timer = QTimer()
        self._slider_save_timer.setSingleShot(True)
        self._slider_save_timer.timeout.connect(self._save_pending_sliders)
        self._pending_slider_changes: dict[str, float] = {}

        self._build_menu_bar()
        self._build_toolbar()
        self._build_central()
        self._build_statusbar()
        self._load_stacks()

        # Wire up collected widgets
        self._wire_voices_and_sliders()

        # Reactive subscriptions — this is what structurally prevents the
        # "UI shows stale/wrong state" bug class: every widget update below
        # is driven by an AppState signal, regardless of what triggered the
        # underlying change (this window's own combo, the tray menu, Load
        # Settings, ...), instead of only updating in response to this
        # window's own widget callbacks.
        self._state.stack_changed.connect(self._on_state_stack_changed)
        self._state.model_changed.connect(self._on_state_model_changed)
        self._state.voice_changed.connect(self._on_state_voice_changed)
        self._state.params_changed.connect(self._on_state_params_changed)
        self._state.speaking_changed.connect(self._on_state_speaking_changed)
        self._state.error_changed.connect(self._on_state_error_changed)
        self._state.catalog_changed.connect(self._on_state_catalog_changed)

    # ── Menu bar ─────────────────────────────────────────────────────────────

    def _build_menu_bar(self) -> None:
        mb = QMenuBar()
        settings_menu = mb.addMenu("&Settings")

        save_action = QAction("&Save Settings…", self)
        save_action.setToolTip("Export current settings to a YAML file")
        save_action.triggered.connect(self._on_save_settings)
        settings_menu.addAction(save_action)

        load_action = QAction("&Load Settings…", self)
        load_action.setToolTip("Import settings from a YAML file")
        load_action.triggered.connect(self._on_load_settings)
        settings_menu.addAction(load_action)

        self.setMenuBar(mb)

    def _on_save_settings(self) -> None:
        path, _ = QFileDialog.getSaveFileName(
            self, "Save Settings", "alienvox-settings.yaml", "YAML Files (*.yaml *.yml)"
        )
        if not path:
            return
        try:
            self._controller.save_settings_to(Path(path))
            self._set_status(f"Settings saved to {Path(path).name}")
        except Exception as exc:
            self._set_status(f"Save settings failed: {exc}")

    def _on_load_settings(self) -> None:
        path, _ = QFileDialog.getOpenFileName(
            self, "Load Settings", "", "YAML Files (*.yaml *.yml)"
        )
        if not path:
            return
        try:
            self._controller.load_settings_from(Path(path))
            self._set_status(f"Settings loaded from {Path(path).name}")
        except Exception as exc:
            self._set_status(f"Load settings failed: {exc}")

    # ── Toolbar ───────────────────────────────────────────────────────────────

    def _build_toolbar(self) -> None:
        tb = QToolBar()
        tb.setMovable(False)
        tb.setFloatable(False)
        tb.setIconSize(QSize(16, 16))
        tb.setStyleSheet(f"""
            QToolBar {{
                background: {_TOOLBAR_BG};
                border-bottom: 1px solid #d0d0d0;
                spacing: 2px;
                padding: 2px 4px;
            }}
            QToolButton {{
                border: 1px solid transparent;
                background: transparent;
                padding: 3px 5px;
                font-size: 13px;
                min-width: 24px;
                min-height: 24px;
            }}
            QToolButton:hover {{
                border-color: #a0a0a0;
                background: #e8e8e8;
            }}
            QToolButton:disabled {{ color: #9a9a9a; }}
        """)

        def _text_btn(text: str, tip: str, slot=None) -> QToolButton:
            b = QToolButton()
            b.setText(text)
            b.setToolTip(tip)
            if slot:
                b.clicked.connect(slot)
            tb.addWidget(b)
            return b

        def _icon_btn(icon: QIcon, tip: str, slot=None) -> QToolButton:
            b = QToolButton()
            b.setIcon(icon)
            b.setIconSize(QSize(16, 16))
            b.setToolTip(tip)
            if slot:
                b.clicked.connect(slot)
            tb.addWidget(b)
            return b

        def _sep() -> None:
            sep = QFrame()
            sep.setFrameShape(QFrame.Shape.VLine)
            sep.setStyleSheet("color: #d0d0d0; margin: 3px 2px;")
            tb.addWidget(sep)

        _text_btn("📄", "New document",  self._on_new)
        _text_btn("📂", "Open file (TXT / MD / HTML / DOCX / PDF)", self._on_open)
        _text_btn("💾", "Save document", self._on_save)
        _sep()
        _text_btn("🎵", "Export to WAV / MP3", self._on_export)
        _sep()
        self._btn_play  = _icon_btn(_make_play_icon(),  "Play (speak text)", self._on_play)
        self._btn_pause = _icon_btn(_make_pause_icon(), "Pause",             self._on_pause)
        self._btn_stop  = _icon_btn(_make_stop_icon(),  "Stop",              self._on_stop)

        # Spacer to push About (and the GPU indicator) to the right
        spacer = QWidget()
        spacer.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        tb.addWidget(spacer)

        if cuda_available():
            _sep()
            gpu_lbl = QLabel()
            gpu_icon_path = Path(__file__).parent / "resources" / "icons" / "gpu.png"
            if gpu_icon_path.exists():
                gpu_lbl.setPixmap(QPixmap(str(gpu_icon_path)).scaled(
                    16, 16, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation,
                ))
            gpu_lbl.setToolTip("Running on GPU (CUDA) — pass --cpu to run.py to force CPU-only")
            tb.addWidget(gpu_lbl)

        _sep()
        _icon_btn(_make_about_icon(), "About AlienVox", self._on_about)
        self.addToolBar(tb)

    # ── Central widget ────────────────────────────────────────────────────────

    def _build_central(self) -> None:
        central = QWidget()
        self.setCentralWidget(central)
        vbox = QVBoxLayout(central)
        vbox.setContentsMargins(0, 0, 0, 0)
        vbox.setSpacing(0)

        self._tabs = QTabWidget()
        self._tabs.setTabPosition(QTabWidget.TabPosition.North)
        self._tabs.setDocumentMode(True)
        self._tabs.setStyleSheet(f"""
            QTabBar::tab {{
                padding: 6px 14px;
                font-size: 12px;
                background: {_TABS_BG};
                border: none;
                border-bottom: 2px solid transparent;
            }}
            QTabBar::tab:selected {{
                background: #ffffff;
                border-bottom: 2px solid {_ACCENT};
                font-weight: 500;
            }}
            QTabBar::tab:hover:!selected {{ background:#d8d8d8; }}
            QTabBar::tab:disabled {{ color:#aaa; }}
            QTabWidget::pane {{ border:none; border-top:1px solid #d0d0d0; }}
        """)
        self._tabs.currentChanged.connect(self._on_tab_changed)
        vbox.addWidget(self._tabs)

        # Shared editor — lives below the engine tabs so content is never
        # lost when switching stacks.
        self._editor = _MultiFormatEditor()
        self._editor.textChanged.connect(
            lambda: self._chars_lbl.setText(
                f"{len(self._editor.to_plain_text())} chars"
            )
        )
        vbox.addWidget(self._editor, stretch=1)

    def _build_statusbar(self) -> None:
        sb = QStatusBar()
        sb.setStyleSheet(
            "QStatusBar { background:#f0f0f0; border-top:1px solid #d0d0d0; "
            "font-size:11px; color:#666; }"
        )
        self._status_lbl = QLabel("Ready")
        self._chars_lbl  = QLabel("0 chars")
        sb.addWidget(self._status_lbl)
        sb.addPermanentWidget(self._chars_lbl)
        self.setStatusBar(sb)

    # ── AppState reactive handlers ───────────────────────────────────────────
    # Every one of these is the ONLY place its corresponding widget state
    # gets updated — no constructor snapshot, no separate "push" method
    # called ad hoc from outside. Whatever triggered the AppState change
    # (this window, the tray, Load Settings) ends up here the same way.

    def _on_state_stack_changed(self, stack_id: str) -> None:
        for i in range(self._tabs.count()):
            tab = self._tabs.widget(i)
            if tab and tab.property("stack_id") == stack_id:
                self._tabs.blockSignals(True)
                self._tabs.setCurrentIndex(i)
                self._tabs.blockSignals(False)
                break
        self._set_status(f"Engine: {stack_id or '—'}")

    def _on_state_model_changed(self, model_id: str) -> None:
        stack_id = self._state.active_stack
        stack = self._state.stack_info(stack_id)
        if stack is None:
            return
        model = self._state.model_info(stack_id, model_id)

        for combo, sid in self._model_combos:
            if sid != stack_id:
                continue
            combo.blockSignals(True)
            idx = combo.findData(model_id)
            if idx >= 0:
                combo.setCurrentIndex(idx)
            combo.blockSignals(False)

        for combo, sid in self._voice_combos:
            if sid != stack_id:
                continue
            combo.blockSignals(True)
            combo.clear()
            if model:
                for v in model.voices:
                    combo.addItem(v.get("label", v["id"]), v["id"])
            combo.blockSignals(False)

    def _on_state_voice_changed(self, voice_id: str) -> None:
        stack_id = self._state.active_stack
        for combo, sid in self._voice_combos:
            if sid != stack_id:
                continue
            combo.blockSignals(True)
            idx = combo.findData(voice_id)
            if idx >= 0:
                combo.setCurrentIndex(idx)
            combo.blockSignals(False)

    def _on_state_params_changed(self, changed: dict) -> None:
        # rate/pitch/volume are global (one value shared across every
        # stack tab's slider row, not per-model) — update every slider
        # with a matching name, not just the active tab's.
        for name, slider in self._sliders:
            if name in changed:
                slider.blockSignals(True)
                slider.setValue(changed[name])
                slider.blockSignals(False)

    def _on_state_speaking_changed(self, speaking: bool) -> None:
        self._set_status("Speaking…" if speaking else "Ready")

    def _on_state_error_changed(self, message: str) -> None:
        if message:
            self._set_status(f"Error: {message}")

    def _on_state_catalog_changed(self) -> None:
        """Live voices (e.g. SAPI's OS-enumerated list) arrived/changed —
        refresh any non-ML voice combo that sources from them."""
        for combo, sid in self._voice_combos:
            stack = self._state.stack_info(sid)
            if stack is None or stack.models:
                continue  # ML stacks source voices from the catalog, not live_voices
            live = self._state.live_voices_for(sid)
            if not live:
                continue
            current_voice = self._state.voice if sid == self._state.active_stack else ""
            self._fill_voice_combo(combo, live, current_voice)

    def _fill_voice_combo(
        self, combo: QComboBox, voices: list[dict], current_voice_id: str = ""
    ) -> None:
        combo.blockSignals(True)
        combo.clear()
        select_idx = 0
        for i, v in enumerate(voices):
            combo.addItem(v.get("label", v["id"]), v["id"])
            if v["id"] == current_voice_id:
                select_idx = i
        if current_voice_id:
            combo.setCurrentIndex(select_idx)
        combo.blockSignals(False)

    # ── User-input wiring ────────────────────────────────────────────────────

    def _wire_voices_and_sliders(self) -> None:
        """Connect voice combos and sliders after all tabs are built."""
        for combo, stack_id in self._voice_combos:
            combo.currentIndexChanged.connect(
                lambda idx, c=combo, s=stack_id: self._on_voice_combo_changed(c, s)
            )

        for name, slider in self._sliders:
            slider.valueChanged.connect(
                lambda v, n=name: self._on_slider_debounced(n, v)
            )

    def _on_voice_combo_changed(self, combo: QComboBox, stack_id: str) -> None:
        """User picked a different voice in a stack's dropdown."""
        vid = combo.itemData(combo.currentIndex())
        if not vid:
            return
        if stack_id == self._state.active_stack:
            self._controller.select_voice(vid)
        else:
            # Defensive: shouldn't normally happen (picking a voice under a
            # non-active tab implies a tab switch, which _on_tab_changed
            # already handles before this could fire) — but if it does,
            # treat it as switching to that stack, not a silent voice-only
            # update against the wrong active model.
            self._controller.select_stack(stack_id, vid)

    def _on_slider_debounced(self, name: str, value: float) -> None:
        """Called on every slider change; debounces saves via QTimer."""
        self._pending_slider_changes[name] = value
        self._slider_save_timer.start(350)  # 350ms debounce

    def _save_pending_sliders(self) -> None:
        """Push pending slider changes to AppController and clear the pending dict."""
        if not self._pending_slider_changes:
            return
        patch = {k: int(v) for k, v in self._pending_slider_changes.items()}
        self._controller.update_params(**patch)
        self._pending_slider_changes.clear()

    # ── Stack tabs ────────────────────────────────────────────────────────────

    def _load_stacks(self) -> None:
        # SAPI 4 — always disabled (legacy)
        dummy = QWidget()
        self._tabs.addTab(dummy, "SAPI 4")
        self._tabs.setTabEnabled(0, False)

        for stack in self._state.stacks:
            tab = self._build_stack_tab(stack)
            self._tabs.addTab(tab, stack.name)
            if not stack.available:
                idx = self._tabs.count() - 1
                self._tabs.setTabEnabled(idx, False)

        # Restore last active stack, falling back to first available tab
        restored = False
        active_stack_id = self._state.active_stack
        if active_stack_id:
            for i in range(self._tabs.count()):
                tab = self._tabs.widget(i)
                if (
                    tab
                    and tab.property("stack_id") == active_stack_id
                    and self._tabs.isTabEnabled(i)
                ):
                    self._tabs.setCurrentIndex(i)
                    restored = True
                    break
        if not restored:
            for i in range(self._tabs.count()):
                if self._tabs.isTabEnabled(i):
                    self._tabs.setCurrentIndex(i)
                    break

    def _build_stack_tab(self, stack: StackInfo) -> QWidget:
        w = QWidget()
        w.setProperty("stack_id", stack.id)
        vbox = QVBoxLayout(w)
        vbox.setContentsMargins(0, 0, 0, 0)
        vbox.setSpacing(0)

        if not stack.available:
            unavail = QLabel(f"<i>{stack.name} — {stack.platform_reason or 'not available on this platform'}</i>")
            unavail.setAlignment(Qt.AlignmentFlag.AlignCenter)
            vbox.addWidget(unavail, stretch=1)
            return w

        # Voice controls bar
        vbox.addWidget(self._build_voice_bar(stack))
        vbox.addWidget(self._hsep())

        # Audio sliders
        vbox.addWidget(self._build_slider_strip(stack))
        vbox.addWidget(self._hsep())

        return w

    def _build_voice_bar(self, stack: StackInfo) -> QWidget:
        bar = QWidget()
        bar.setStyleSheet(f"background:{_VOICEBAR_BG}; border-bottom:1px solid #d0d0d0;")
        layout = QHBoxLayout(bar)
        layout.setContentsMargins(8, 6, 8, 6)
        layout.setSpacing(8)

        has_models = bool(stack.models)

        # Which model is actually active for this stack right now — sourced
        # from AppState, not stack.models[0]. Only the actually active
        # stack's tab tries to match AppState's current model/voice;
        # non-active tabs just show their first model/voice, since
        # AppState's values weren't recorded for a stack that isn't running.
        is_active_stack = stack.id == self._state.active_stack
        selected_model = None
        if has_models:
            selected_model = (
                next((m for m in stack.models if m.id == self._state.active_model), None)
                if is_active_stack else None
            ) or stack.models[0]

        if has_models:
            model_combo = QComboBox()
            model_combo.setFixedWidth(220)
            model_combo.setStyleSheet(_combo_style())
            for i, m in enumerate(stack.models):
                model_combo.addItem(m.name, m.id)
                if m.id == selected_model.id:
                    model_combo.setCurrentIndex(i)
            bar.setProperty("model_combo", model_combo)
            self._model_combos.append((model_combo, stack.id))
            layout.addWidget(model_combo)

        voice_combo = QComboBox()
        voice_combo.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        voice_combo.setStyleSheet(_combo_style())
        bar.setProperty("voice_combo", voice_combo)
        self._voice_combos.append((voice_combo, stack.id))

        if has_models and selected_model is not None:
            select_idx = 0
            for i, v in enumerate(selected_model.voices):
                voice_combo.addItem(v.get("label", v["id"]), v["id"])
                if is_active_stack and v["id"] == self._state.voice:
                    select_idx = i
            if is_active_stack and self._state.voice:
                voice_combo.setCurrentIndex(select_idx)
        else:
            live = self._state.live_voices_for(stack.id)
            if live:
                self._fill_voice_combo(voice_combo, live, self._state.voice if is_active_stack else "")
            elif stack.id in ("sapi5", "speech_platform"):
                voice_combo.addItem("(loading voices…)", "")
            else:
                voices = get_voices(stack.id)
                for v in voices:
                    voice_combo.addItem(v.get("label", v["id"]), v["id"])

        layout.addWidget(voice_combo)

        if has_models:
            # TTL spinbox
            from PySide6.QtWidgets import QSpinBox
            ttl_lbl = QLabel("TTL")
            ttl_lbl.setStyleSheet("font-size:11px; color:#555;")
            ttl_spin = QSpinBox()
            ttl_spin.setRange(0, 300)
            ttl_spin.setValue(self._state.ttl_seconds)
            ttl_spin.setSuffix(" s")
            ttl_spin.setFixedWidth(72)
            ttl_spin.setStyleSheet(_combo_style())
            layout.addWidget(ttl_lbl)
            layout.addWidget(ttl_spin)

        self._status_engine_lbl = QLabel("Ready")
        self._status_engine_lbl.setStyleSheet("font-size:11px; color:#555; min-width:100px;")
        layout.addWidget(self._status_engine_lbl)

        install_btn = QPushButton("Install" if not has_models else "Install Model")
        install_btn.setStyleSheet(_btn_style())
        install_btn.setFixedHeight(28)
        install_btn.clicked.connect(lambda: self._open_install_dialog(stack))
        layout.addWidget(install_btn)

        if has_models:
            model_combo.currentIndexChanged.connect(
                lambda idx, s=stack, mc=model_combo:
                    self._on_model_combo_changed(s, mc.itemData(idx))
            )

        return bar

    def _build_slider_strip(self, stack: StackInfo) -> QWidget:
        strip = QWidget()
        strip.setStyleSheet(f"background:{_SLIDER_BG};")
        grid = QHBoxLayout(strip)
        grid.setContentsMargins(8, 10, 8, 10)
        grid.setSpacing(20)

        cfg = load_effective_config(stack.id, user_file=None)
        controls = get_controls(stack.id)

        rate_spec   = controls.get("rate",   {})
        pitch_spec  = controls.get("pitch",  {})
        vol_spec    = controls.get("volume", {})

        s_rate = _SliderRow(
            "Rate", float(rate_spec.get("min", -10)), float(rate_spec.get("max", 10)),
            float(cfg.get("rate", rate_spec.get("default", 0))),
            enabled=rate_spec.get("applies", True),
        )
        s_pitch = _SliderRow(
            "Pitch", float(pitch_spec.get("min", -10)), float(pitch_spec.get("max", 10)),
            float(cfg.get("pitch", pitch_spec.get("default", 0))),
            enabled=pitch_spec.get("applies", True),
        )
        if not pitch_spec.get("applies", True):
            s_pitch._name_lbl.setText("Pitch (N/A)")

        s_vol = _SliderRow(
            "Volume", float(vol_spec.get("min", 0)), float(vol_spec.get("max", 100)),
            float(cfg.get("volume", vol_spec.get("default", 100))),
            enabled=vol_spec.get("applies", True),
        )

        for name, s in (("rate", s_rate), ("pitch", s_pitch), ("volume", s_vol)):
            grid.addWidget(s)
            self._sliders.append((name, s))

        strip.setProperty("slider_rate",   s_rate)
        strip.setProperty("slider_pitch",  s_pitch)
        strip.setProperty("slider_volume", s_vol)

        return strip

    # ── Toolbar actions ───────────────────────────────────────────────────────

    def _on_new(self) -> None:
        self._editor.clear()
        self._set_status("Ready")

    def _on_open(self) -> None:
        start = str(self._editor.current_path().parent) if self._editor.current_path() else ""
        path = self._editor.open_file(start_dir=start)
        if path:
            fmt = self._editor.current_fmt().upper()
            self._set_status(f"Opened {path.name} [{fmt}]")

    def _on_save(self) -> None:
        path = self._editor.save_file()
        if path:
            self._set_status(f"Saved {path.name}")

    def _on_export(self) -> None:
        text = self._editor.to_plain_text().strip()
        if not text:
            self._set_status("No text to export")
            return
        if not self._controller.engine:
            self._set_status("Export not available — engine not loaded")
            return
        from .export_dialog import ExportDialog
        dlg = ExportDialog(
            engine=self._controller.engine,
            text=text,
            voice_id=self._state.voice,
            params=self._controller.build_current_speak_params(),
            parent=self,
        )
        dlg.exec()

    def _on_play(self) -> None:
        text = self._editor.to_plain_text().strip()
        if not text:
            self._set_status("No text to speak")
            return
        self._set_status("Speaking…")
        self._controller.play_async(text)

    def _on_pause(self) -> None:
        self._set_status("Paused")  # pause/resume via engine — wired later

    def _on_stop(self) -> None:
        self._controller.stop()
        self._set_status("Stopped")

    def _on_about(self) -> None:
        from .about import AboutDialog
        dlg = AboutDialog(parent=self)
        dlg.exec()

    def _on_tab_changed(self, idx: int) -> None:
        tab = self._tabs.widget(idx)
        if not tab:
            return
        stack_id = tab.property("stack_id") or ""
        if not stack_id or stack_id == self._state.active_stack:
            return
        voice_id = ""
        for combo, sid in self._voice_combos:
            if sid == stack_id:
                voice_id = combo.currentData() or ""
                break
        self._controller.select_stack(stack_id, voice_id)

    def _open_install_dialog(self, stack: StackInfo) -> None:
        from ..config import models_root as _models_root
        from .install_dialog import InstallDialog

        # Resolve the currently selected model for this stack
        active_model_id = ""
        if stack.models:
            active_model_id = stack.models[0].id
            for combo, sid in getattr(self, "_model_combos", []):
                if sid == stack.id:
                    active_model_id = combo.currentData() or active_model_id
                    break

        dlg = InstallDialog(
            stack=stack,
            models_root=self._models_root or _models_root(),
            active_model_id=active_model_id,
            parent=self,
        )
        dlg.exec()

    def _on_model_combo_changed(self, stack: StackInfo, model_id: str) -> None:
        """User picked a different model in the ML/AI tab's dropdown.

        Just forwards to AppController — the actual voice-combo repopulation
        happens reactively in _on_state_model_changed once AppState's
        model_changed signal fires, whether it came from this combo, the
        tray menu, or Load Settings.
        """
        if not model_id:
            return
        model = next((m for m in stack.models if m.id == model_id), None)
        first_voice_id = model.voices[0]["id"] if model and model.voices else ""
        self._controller.select_model(model_id, first_voice_id)

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _active_editor(self) -> _MultiFormatEditor:
        return self._editor

    def _set_status(self, msg: str) -> None:
        self._status_lbl.setText(msg)
        if hasattr(self, "_status_engine_lbl"):
            self._status_engine_lbl.setText(msg)

    @staticmethod
    def _hsep() -> QFrame:
        f = QFrame()
        f.setFrameShape(QFrame.Shape.HLine)
        f.setStyleSheet("color:#d0d0d0;")
        return f


# ── Toolbar icon painters ─────────────────────────────────────────────────────
# Painted icons so Qt doesn't render emoji glyphs (which ignore CSS color).
# Sizes and colours match the Rust/Tauri CSS:
#   play  → #22c55e (green)   pause → #6b7280 (gray)   stop → #ef4444 (red)

def _make_icon(size: int = 16) -> QPixmap:
    pix = QPixmap(size, size)
    pix.fill(Qt.GlobalColor.transparent)
    return pix


def _make_play_icon(size: int = 16) -> QIcon:
    pix = _make_icon(size)
    p = QPainter(pix)
    p.setRenderHint(QPainter.RenderHint.Antialiasing)
    p.setBrush(QColor("#22c55e"))
    p.setPen(Qt.PenStyle.NoPen)
    m = size // 5
    tri = QPolygon([
        QPoint(m, m),
        QPoint(size - m, size // 2),
        QPoint(m, size - m),
    ])
    p.drawPolygon(tri)
    p.end()
    return QIcon(pix)


def _make_pause_icon(size: int = 16) -> QIcon:
    pix = _make_icon(size)
    p = QPainter(pix)
    p.setRenderHint(QPainter.RenderHint.Antialiasing)
    p.setBrush(QColor("#6b7280"))
    p.setPen(Qt.PenStyle.NoPen)
    bar_w = max(3, size // 4)
    gap   = max(2, size // 5)
    x1 = (size - bar_w * 2 - gap) // 2
    x2 = x1 + bar_w + gap
    my = size // 5
    p.drawRect(x1, my, bar_w, size - my * 2)
    p.drawRect(x2, my, bar_w, size - my * 2)
    p.end()
    return QIcon(pix)


def _make_stop_icon(size: int = 16) -> QIcon:
    pix = _make_icon(size)
    p = QPainter(pix)
    p.setRenderHint(QPainter.RenderHint.Antialiasing)
    p.setBrush(QColor("#ef4444"))
    p.setPen(Qt.PenStyle.NoPen)
    m = size // 4
    p.drawRect(m, m, size - m * 2, size - m * 2)
    p.end()
    return QIcon(pix)


def _make_about_icon(size: int = 16) -> QIcon:
    """App logo scaled to toolbar size — used for the About button."""
    icons_dir = Path(__file__).parent / "resources" / "icons"
    logo = icons_dir / "icon_16x16.png"
    if logo.exists():
        pix = QPixmap(str(logo))
        if not pix.isNull():
            return QIcon(pix)
    # Fallback: draw a blue circle with a white "i"
    pix = _make_icon(size)
    p = QPainter(pix)
    p.setRenderHint(QPainter.RenderHint.Antialiasing)
    p.setBrush(QColor("#0078d4"))
    p.setPen(Qt.PenStyle.NoPen)
    p.drawEllipse(1, 1, size - 2, size - 2)
    p.setPen(QColor("#ffffff"))
    f = p.font()
    f.setBold(True)
    f.setPixelSize(size - 4)
    p.setFont(f)
    p.drawText(pix.rect(), Qt.AlignmentFlag.AlignCenter, "i")
    p.end()
    return QIcon(pix)


# ── Style helpers ─────────────────────────────────────────────────────────────

def _combo_style() -> str:
    return (
        "QComboBox { border:1px solid #c0c0c0; background:#ffffff; "
        "padding:4px 8px; font-size:12px; height:28px; }"
        "QComboBox::drop-down { border:none; }"
        "QComboBox QAbstractItemView { border:1px solid #c0c0c0; font-size:12px; }"
    )


def _btn_style() -> str:
    return (
        "QPushButton { border:1px solid #c0c0c0; background:#ffffff; "
        "padding:4px 14px; font-size:12px; }"
        "QPushButton:hover { background:#e8e8e8; }"
    )
